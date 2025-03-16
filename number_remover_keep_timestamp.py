from docx import Document
import os
import re

def is_timestamp_line(text):
    """
    Determine if a line is a timestamp.
    Works with multiple timestamp formats.
    """
    # Check for timestamp patterns
    
    # Format: 00:00:02.895 --> 00:00:03.765
    long_timestamp = re.match(r'^\d{2}:\d{2}:\d{2}\.\d{3}\s*-->\s*\d{2}:\d{2}:\d{2}\.\d{3}$', text.strip())
    
    # Format: 0:01, 1:42, etc.
    short_timestamp = re.match(r'^\d+:\d{2}$', text.strip())
    
    return bool(long_timestamp or short_timestamp)

def is_standalone_number(text, prev_line, next_line):
    """
    Determine if a line is just a number that should be removed.
    Takes into account the context (previous and next lines).
    """
    # Check if the line is just a number
    if not re.match(r'^\s*\d+\s*$', text.strip()):
        return False
    
    # If previous line is a timestamp, this is likely an auto-generated number
    if prev_line and is_timestamp_line(prev_line):
        return True
    
    # If next line is a timestamp, this might be an auto-generated number
    if next_line and is_timestamp_line(next_line):
        return True
    
    # Additional check: if line is a single number at the start of a section
    # followed by content (not a timestamp), it's likely an auto-generated number
    if (not prev_line or prev_line.strip() == '') and next_line and not is_timestamp_line(next_line):
        return True
    
    return False

def clean_transcript(input_file, output_file=None, keep_timestamps=True):
    """
    Clean a transcript by removing auto-generated numbers, with option to keep timestamps.
    
    Args:
        input_file (str): Path to the input DOCX file
        output_file (str, optional): Path for the output file. If None, creates a file with _cleaned suffix
        keep_timestamps (bool): Whether to preserve timestamp lines in the output
        
    Returns:
        str: Path to the cleaned file
    """
    # Set output filename if not provided
    if not output_file:
        file_base, file_ext = os.path.splitext(input_file)
        output_file = f"{file_base}_cleaned{file_ext}"
    
    try:
        # Load the document
        doc = Document(input_file)
        
        # Extract all paragraphs with text
        paragraphs = [p.text for p in doc.paragraphs]
        
        # Create a new document for cleaned content
        new_doc = Document()
        
        # Process each paragraph with context awareness
        for i, text in enumerate(paragraphs):
            # Skip empty paragraphs
            if not text.strip():
                continue
                
            # Get previous and next lines for context
            prev_line = paragraphs[i-1] if i > 0 else None
            next_line = paragraphs[i+1] if i < len(paragraphs)-1 else None
            
            # Handle timestamps according to preference
            if is_timestamp_line(text):
                if keep_timestamps:
                    new_doc.add_paragraph(text)
                continue
                
            # Skip standalone numbers we want to remove
            if is_standalone_number(text, prev_line, next_line):
                continue
                
            # Add paragraph to new document
            new_doc.add_paragraph(text)
        
        # Save the cleaned document
        new_doc.save(output_file)
        print(f"Cleaned transcript saved to: {output_file}")
        return output_file
        
    except Exception as e:
        print(f"Error processing file {input_file}: {str(e)}")
        return None

def process_directory(input_dir, output_dir=None, keep_timestamps=True):
    """
    Process all DOCX files in a directory.
    
    Args:
        input_dir (str): Directory containing files to clean
        output_dir (str, optional): Directory to save cleaned files
        keep_timestamps (bool): Whether to preserve timestamp lines
        
    Returns:
        list: Paths to all cleaned files
    """
    # Set output directory
    if not output_dir:
        output_dir = input_dir
        
    # Create output directory if it doesn't exist
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    cleaned_files = []
    
    # Process each DOCX file
    for filename in os.listdir(input_dir):
        if filename.endswith('.docx'):
            input_path = os.path.join(input_dir, filename)
            base_name, ext = os.path.splitext(filename)
            output_path = os.path.join(output_dir, f"{base_name}_cleaned{ext}")
            
            result = clean_transcript(input_path, output_path, keep_timestamps)
            if result:
                cleaned_files.append(result)
    
    return cleaned_files

if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="Clean transcript files by removing auto-generated numbers")
    parser.add_argument("input", help="Input file or directory")
    parser.add_argument("--output", "-o", help="Output file or directory (optional)")
    parser.add_argument("--batch", "-b", action="store_true", help="Process directory instead of single file")
    parser.add_argument("--remove-timestamps", "-r", action="store_true", 
                        help="Remove timestamp lines as well as numbering (default: keep timestamps)")
    
    args = parser.parse_args()
    keep_timestamps = not args.remove_timestamps
    
    if args.batch:
        results = process_directory(args.input, args.output, keep_timestamps)
        print(f"Processed {len(results)} files")
    else:
        clean_transcript(args.input, args.output, keep_timestamps)