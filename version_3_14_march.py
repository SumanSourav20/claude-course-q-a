import os
import re
import math
import docx
import anthropic
import argparse
from typing import List, Dict, Tuple, Optional, Any

explanation = """
	explanation: 
		A: why it's right or wrong,
		B: why it's right or wrong,
		C: why it's right or wrong,
		D: why it's right or wrong,	
"""

class TranscriptProcessor:
    def __init__(self, api_key: str, model: str = "claude-3-7-sonnet-20250219"):
        """
        Initialize the transcript processor with API key and model.
        
        Args:
            api_key: Anthropic API key
            model: Claude model to use
        """
        self.client = anthropic.Anthropic(api_key=api_key)
        self.model = model
        self.summaries = []
        self.questions = []

    def extract_text_from_docx(self, file_path: str) -> Tuple[str, int]:
        """
        Extract text from a docx file including timestamps and count pages.
        
        Args:
            file_path: Path to the docx file
            
        Returns:
            Tuple of (full_transcript_text, page_count)
        """
        doc = docx.Document(file_path)
        full_text = []
        
        # Try to get actual page count from document properties
        try:
            # This works for some docx files that store page count in document properties
            actual_page_count = doc.core_properties.page
            if actual_page_count and actual_page_count > 0:
                print(f"Document reports {actual_page_count} pages")
                return "\n".join([para.text for para in doc.paragraphs]), actual_page_count
        except:
            pass
            
        # Collect text and estimate based on multiple methods
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        full_text_joined = "\n".join(full_text)
            
        # Method 1: Estimate based on word count (commonly used but not perfect)
        words = full_text_joined.split()
        word_count = len(words)
        words_per_page_estimate = 250  # Standard estimate for typical documents
        est_pages_by_words = max(1, word_count // words_per_page_estimate)
        
        # Method 2: Estimate based on paragraph count
        paragraphs = [p for p in full_text if p.strip()]
        paragraphs_per_page_estimate = 15  # Standard estimate
        est_pages_by_paragraphs = max(1, len(paragraphs) // paragraphs_per_page_estimate)
        
        # Method 3: Estimate based on character count
        chars = len(full_text_joined)
        chars_per_page_estimate = 2000  # Standard estimate including spaces
        est_pages_by_chars = max(1, chars // chars_per_page_estimate)
        
        # Special handling for transcript format (lots of timestamps)
        # If transcript format detected, adjust estimates
        timestamp_pattern = r'\d{1,2}:\d{2}:\d{2}|\d{1,2}:\d{2}'
        timestamps = re.findall(timestamp_pattern, full_text_joined)
        if len(timestamps) > 50:  # Likely a transcript with many timestamps
            # Transcripts usually have fewer words per page due to formatting
            transcript_adjustment = 1.75  # Adjustment factor for transcripts
            est_pages_by_words = int(est_pages_by_words * transcript_adjustment)
            est_pages_by_chars = int(est_pages_by_chars * transcript_adjustment)
        
        # Use a weighted average of the estimates with more weight on word count
        estimated_pages = int((est_pages_by_words * 0.5) + 
                            (est_pages_by_paragraphs * 0.2) + 
                            (est_pages_by_chars * 0.3))
        
        # Allow user to override the estimate with an environment variable
        if 'TRANSCRIPT_PAGE_COUNT' in os.environ:
            try:
                user_specified_pages = int(os.environ['TRANSCRIPT_PAGE_COUNT'])
                print(f"Using user-specified page count: {user_specified_pages}")
                return full_text_joined, user_specified_pages
            except:
                pass
        
        print(f"Estimated page count: {estimated_pages} pages (based on {word_count} words, {len(paragraphs)} paragraphs, {chars} characters)")
        print(f"Individual estimates: By words: {est_pages_by_words}, By paragraphs: {est_pages_by_paragraphs}, By chars: {est_pages_by_chars}")
        
        return full_text_joined, estimated_pages

    def chunk_transcript(self, transcript: str, total_pages: int, target_pages_per_chunk: int = 30) -> List[str]:
        """
        Divide the transcript into chunks based on page count.
        
        Args:
            transcript: Full transcript text
            total_pages: Estimated total number of pages
            target_pages_per_chunk: Target number of pages per chunk (default: 30)
            
        Returns:
            List of chunked transcript sections
        """
        lines = transcript.split('\n')
        # Remove empty lines
        lines = [line for line in lines if line.strip()]
        
        # Determine optimal number of chunks
        if total_pages <= target_pages_per_chunk + 5:  # If close to target, keep as single chunk
            # If it's already reasonably close to our target (30-35 pages), keep as one chunk
            return [transcript]
            
        # Calculate optimal number of chunks to keep each close to target size
        num_chunks = max(1, round(total_pages / target_pages_per_chunk))
        
        # If dividing would make chunks too small (< 15 pages), reduce number of chunks
        pages_per_chunk = total_pages / num_chunks
        if pages_per_chunk < 15 and num_chunks > 1:
            num_chunks -= 1
            pages_per_chunk = total_pages / num_chunks
        
        # Calculate lines per chunk based on estimated pages
        lines_per_chunk = math.ceil(len(lines) / num_chunks)
        
        chunks = []
        for i in range(0, len(lines), lines_per_chunk):
            chunk = '\n'.join(lines[i:i + lines_per_chunk])
            chunks.append(chunk)
        
        print(f"Divided {total_pages} pages into {len(chunks)} chunks " + 
              f"(~{total_pages/len(chunks):.1f} pages per chunk)")
            
        return chunks
    
    def extract_timestamp_range(self, chunk: str) -> Tuple[str, str]:
        """
        Extract the first and last timestamp from a chunk.
        
        Args:
            chunk: Transcript chunk
            
        Returns:
            Tuple of (first_timestamp, last_timestamp)
        """
        # Common timestamp patterns (adjust as needed for your transcript format)
        timestamp_pattern = r'(\d{1,2}:\d{2}:\d{2}|\d{1,2}:\d{2})'
        
        timestamps = re.findall(timestamp_pattern, chunk)
        
        if timestamps:
            return timestamps[0], timestamps[-1]
        return "00:00", "00:00"  # Default if no timestamps found

    def generate_first_chunk_prompt(self, chunk: str, total_chunks: int) -> str:
        """
        Generate prompt for the first chunk, instructing to avoid irrelevant content.
        
        Args:
            chunk: First transcript chunk
            total_chunks: Total number of chunks
            
        Returns:
            Formatted prompt for Claude
        """
        first_timestamp, last_timestamp = self.extract_timestamp_range(chunk)
        
        # If only one chunk exists, generate 10 questions, otherwise 5
        num_questions = 10 if total_chunks == 1 else 5
        
        prompt = f"""You are a subject matter expert who deeply understands this educational topic. You need to create thoughtful quiz questions that test key concepts from this material.

This section covers content from timestamps {first_timestamp} to {last_timestamp}.

IMPORTANT INSTRUCTIONS:
1. IGNORE all content about course logistics, how to pass the course, certificates, or instructor introductions
2. First, thoroughly read and understand the core educational concepts being taught
3. Generate {num_questions} challenging but fair multiple-choice questions that test understanding of important concepts
4. Each question must:
   - Be directly based on the educational content
   - Include timestamp range from where the question and answer is picked up(beside the question)
   - Have 4 answer options with exactly one correct answer
   - Question and answer should be in simple layman's language and should have keywords from the content provided to you
   - Include an explanation of wrong answers as well as right answers
   - Feel connected to the broader learning journey
5. Explanation should be like this:
    {explanation}

CRITICAL: Your questions and explanations must read as if written by a subject matter expert who genuinely understands the topic.Questions and explainations should be like they are final examination questions, Do NOT use phrases like:
- "According to the transcript..."
- "In this course..."
- "The instructor mentions..."
- "As stated in the lecture..."
- "Based on the material provided..."
- "In the content"
- "as explained in or according to or as stated" etc

Instead, explain concepts authoritatively as established knowledge in the field. The questions should feel like they were hand-crafted by someone with expert understanding of the subject.

After the questions, provide a brief 2-3 sentence summary of the key concepts covered.

Here is the content:
{chunk}
"""
        return prompt

    def generate_subsequent_chunk_prompt(self, chunk: str, previous_summaries: str) -> str:
        """
        Generate prompt for subsequent chunks, including previous summaries for context.
        
        Args:
            chunk: Current transcript chunk
            previous_summaries: Combined summaries from previous chunks
            
        Returns:
            Formatted prompt for Claude
        """
        first_timestamp, last_timestamp = self.extract_timestamp_range(chunk)
        
        num_questions = 5  # Default for subsequent chunks
        
        prompt = f"""You are a subject matter expert who deeply understands this educational topic. You need to create thoughtful quiz questions that test key concepts from this material.

This section covers content from timestamps {first_timestamp} to {last_timestamp}.

Context from previous sections:
{previous_summaries}

IMPORTANT INSTRUCTIONS:
1. First, thoroughly read and understand how this section builds on previous knowledge
2. Generate {num_questions} challenging but fair multiple-choice questions that test understanding of NEW concepts in this section
3. Each question must:
   - Be directly based on the educational content 
   - Include timestamp range from where the question and answer is picked up(beside the question)
   - Have 4 answer options with exactly one correct answer
   - Question and answer should be in simple layman's language and should have keywords from the content provided to you
   - Include an explanation of wrong answers as well as right answers
   - Feel connected to the broader learning journey
4. Explanation should be like this:
    {explanation}

CRITICAL: Your questions and explanations must read as if written by a subject matter expert who genuinely understands the entire topic.Questions and explainations should be like they are final examination questions, Do NOT use phrases like:
- "According to the transcript..."
- "In this course..."
- "The instructor mentions..."
- "As stated in the lecture..."
- "Based on the material provided..."
-"In the content"
- "as explained in or according to or as stated" etc

Instead, explain concepts authoritatively as established knowledge in the field. The questions should feel like they were hand-crafted by someone with expert understanding of the subject who is carefully building upon previously established concepts.

After the questions, provide a brief 2-3 sentence summary of the key concepts covered in this section.

Here is the content for this section:
{chunk}
"""
        return prompt

    def generate_questions_from_chunk(self, chunk: str, is_first_chunk: bool, 
                                      previous_summaries: Optional[str] = None, total_chunks: int = 1) -> Dict:
        """
        Generate MCQ questions from a transcript chunk using Claude.
        
        Args:
            chunk: Transcript chunk
            is_first_chunk: Whether this is the first chunk
            previous_summaries: Combined summaries from previous chunks
            total_chunks: Total number of chunks in the document
            
        Returns:
            Dict containing generated questions and summary
        """
        if is_first_chunk:
            prompt = self.generate_first_chunk_prompt(chunk, total_chunks)
        else:
            prompt = self.generate_subsequent_chunk_prompt(chunk, previous_summaries)
            
        response = self.client.messages.create(
            model=self.model,
            max_tokens=4000,
            messages=[
                {"role": "user", "content": prompt}
            ]
        )
        
        return {"response": response.content[0].text}

    def extract_summary_from_response(self, response: str) -> str:
        """
        Extract the summary from Claude's response.
        
        Args:
            response: Claude's full response
            
        Returns:
            Extracted summary
        """
        # Look for common patterns that might indicate a summary section
        summary_patterns = [
            r"Summary:(.*?)(?=\n\n|\Z)",
            r"Key concepts covered:(.*?)(?=\n\n|\Z)",
            r"Brief summary:(.*?)(?=\n\n|\Z)",
            r"Summary of key concepts:(.*?)(?=\n\n|\Z)",
            r"In summary,(.*?)(?=\n\n|\Z)"
        ]
        
        for pattern in summary_patterns:
            match = re.search(pattern, response, re.DOTALL | re.IGNORECASE)
            if match:
                return match.group(1).strip()
        
        # If no clear summary section, check the last few paragraphs
        paragraphs = response.split("\n\n")
        for para in reversed(paragraphs):
            if len(para.split()) > 10 and "question" not in para.lower() and any(term in para.lower() for term in ["cover", "discuss", "introduce", "explore", "focus", "explain", "concept"]):
                return para.strip()
                
        # Fallback: take the last substantial paragraph
        for para in reversed(paragraphs):
            if len(para.split()) > 15 and not any(q in para.lower() for q in ["question", "option", "correct answer", "explanation"]):
                return para.strip()
                
        return "This section covers key educational concepts from the material."

    def process_transcript(self, file_path: str, output_dir: str = "quiz_output", target_pages_per_chunk: int = 30):
        """
        Process the entire transcript, chunk it, and generate questions.
        
        Args:
            file_path: Path to the transcript docx file
            output_dir: Directory to save output files
            target_pages_per_chunk: Target number of pages per chunk
        """
        os.makedirs(output_dir, exist_ok=True)
        
        # Extract transcript and estimate page count
        transcript, total_pages = self.extract_text_from_docx(file_path)
        print(f"Estimated document size: {total_pages} pages")
        
        # Chunk the transcript based on page count
        chunks = self.chunk_transcript(transcript, total_pages, target_pages_per_chunk)
        
        all_questions = []
        combined_summaries = ""
        
        # Process each chunk
        for i, chunk in enumerate(chunks):
            print(f"Processing chunk {i+1}/{len(chunks)}...")
            
            is_first_chunk = (i == 0)
            
            # Generate questions for this chunk
            result = self.generate_questions_from_chunk(
                chunk, 
                is_first_chunk,
                combined_summaries if not is_first_chunk else None,
                total_chunks=len(chunks)
            )
            
            # Extract and save summary for context in next chunks
            summary = self.extract_summary_from_response(result["response"])
            chunk_summary = f"Chunk {i+1}: {summary}"
            self.summaries.append(chunk_summary)
            
            # Update combined summaries for next chunk
            if combined_summaries:
                combined_summaries += f"\n\n{chunk_summary}"
            else:
                combined_summaries = chunk_summary
                
            # Save the questions for this chunk
            all_questions.append(result["response"])
            
            # Write individual chunk result to file
            with open(f"{output_dir}/chunk_{i+1}_questions.txt", "w", encoding="utf-8") as f:
                f.write(result["response"])
                
        # Write all questions to a single file
        with open(f"{output_dir}/all_questions.txt", "w", encoding="utf-8") as f:
            for i, questions in enumerate(all_questions):
                f.write(f"=== CHUNK {i+1} ===\n\n")
                f.write(questions)
                f.write("\n\n")
                
        # Write summaries to a file
        with open(f"{output_dir}/summaries.txt", "w", encoding="utf-8") as f:
            f.write("\n\n".join(self.summaries))
            
        print(f"Processing complete. Results saved to {output_dir}/")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Generate quiz questions from course transcript')
    parser.add_argument('--file', type=str, required=True, help='Path to transcript docx file')
    parser.add_argument('--api-key', type=str, required=True, help='Anthropic API key')
    parser.add_argument('--output-dir', type=str, default='quiz_output', help='Output directory')
    parser.add_argument('--model', type=str, default='claude-3-7-sonnet-20250219', help='Claude model to use')
    parser.add_argument('--target-pages', type=int, default=35, 
                       help='Target number of pages per chunk (default: 35)')
    parser.add_argument('--page-count', type=int, help='Override estimated page count with actual count')
    
    args = parser.parse_args()
    
    # Set environment variable if user specified the page count
    if args.page_count:
        os.environ['TRANSCRIPT_PAGE_COUNT'] = str(args.page_count)
    
    processor = TranscriptProcessor(api_key=args.api_key, model=args.model)
    processor.process_transcript(args.file, args.output_dir, args.target_pages)