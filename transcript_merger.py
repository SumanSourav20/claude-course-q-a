import re
import os
import argparse
from typing import List, Tuple, Optional
import docx
from docx import Document
from docx.shared import Pt

def parse_timestamp(timestamp: str) -> Optional[str]:
    """
    Parse different timestamp formats into a standardized form.
    Handles formats like:
    - 00:11:43.765 --> 00:11:46.385
    - 0:01
    - 0:11 - 0:12
    
    Returns the timestamp in its original format or None if not a timestamp.
    """
    # Format: 00:11:43.765 --> 00:11:46.385
    if re.match(r'\d+:\d+:\d+\.\d+ --> \d+:\d+:\d+\.\d+', timestamp):
        return timestamp
    
    # Format: 0:01
    if re.match(r'^\d+:\d+$', timestamp.strip()):
        return timestamp.strip()
    
    # Format: 0:11 - 0:12
    if re.match(r'^\d+:\d+ - \d+:\d+$', timestamp.strip()):
        return timestamp.strip()
    
    return None

def extract_first_last_timestamp(lines: List[str]) -> Tuple[Optional[str], Optional[str]]:
    """
    Extract the first and last timestamp from a block of text.
    """
    first_timestamp = None
    last_timestamp = None
    
    for line in lines:
        timestamp = parse_timestamp(line)
        if timestamp:
            if first_timestamp is None:
                first_timestamp = timestamp
            last_timestamp = timestamp
    
    return first_timestamp, last_timestamp

def combine_timestamps(first: str, last: str) -> str:
    """
    Combine first and last timestamps into a range.
    """
    # Already a range like '00:11:43.765 --> 00:11:46.385'
    if '-->' in first:
        start = first.split('-->')[0].strip()
        if '-->' in last:
            end = last.split('-->')[1].strip()
        else:
            end = last
        return f"{start} --> {end}"
    
    # Handle simple timestamps like '0:01'
    elif ':' in first and ':' in last and '-' not in first and '-' not in last:
        return f"{first} --> {last}"
    
    # Handle range timestamps like '0:11 - 0:12'
    elif '-' in first and '-' in last:
        start = first.split('-')[0].strip()
        end = last.split('-')[1].strip()
        return f"{start} - {end}"
    
    # Mixed formats - default to simple combination
    else:
        return f"{first} --> {last}"

def read_docx(file_path: str) -> List[str]:
    """
    Read text from a DOCX file and return a list of lines.
    """
    doc = Document(file_path)
    lines = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            lines.append(paragraph.text.strip())
    return lines

def save_to_docx(chunks: List[str], output_file: str):
    """
    Save processed chunks to a new DOCX file.
    """
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    for chunk in chunks:
        # Split into lines - first line is the timestamp, rest is content
        chunk_lines = chunk.strip().split('\n')
        
        if len(chunk_lines) >= 2:
            # Add timestamp as a heading
            heading = doc.add_paragraph(chunk_lines[0])
            heading.style = doc.styles['Heading 2']
            
            # Add content as normal paragraph
            content = doc.add_paragraph(chunk_lines[1])
            
            # Add a blank line between chunks
            doc.add_paragraph()
    
    doc.save(output_file)
    print(f"Saved processed document to {output_file}")

def process_transcript(input_file: str, output_file: str, min_lines: int = 3):
    """
    Process transcript file and combine timestamps for text chunks.
    
    Args:
        input_file: Path to the input transcript file (txt or docx)
        output_file: Path to save the processed output (docx)
        min_lines: Minimum number of lines to consider as a chunk (default: 3)
    """
    # Read input file based on its extension
    if input_file.lower().endswith('.docx'):
        lines = read_docx(input_file)
    else:
        # Assume text file
        with open(input_file, 'r', encoding='utf-8') as f:
            lines = [line.strip() for line in f.readlines() if line.strip()]
    
    chunks = []
    current_chunk = []
    
    for line in lines:
        # If this line is a timestamp and we already have content in the current chunk,
        # it might be the start of a new chunk
        if parse_timestamp(line) and current_chunk and not parse_timestamp(current_chunk[-1]):
            # Check if we have enough lines to consider this a complete chunk
            if len(current_chunk) >= min_lines:
                first_ts, last_ts = extract_first_last_timestamp(current_chunk)
                if first_ts and last_ts:
                    combined_ts = combine_timestamps(first_ts, last_ts)
                    
                    # Filter out timestamp lines and keep only text
                    text_lines = [l for l in current_chunk if not parse_timestamp(l)]
                    text = ' '.join(text_lines)
                    
                    chunks.append(f"{combined_ts}\n{text}")
                
                # Start a new chunk
                current_chunk = [line]
            else:
                # Not enough lines, continue with this chunk
                current_chunk.append(line)
        else:
            # Continue building the current chunk
            current_chunk.append(line)
    
    # Don't forget the last chunk
    if current_chunk and len(current_chunk) >= min_lines:
        first_ts, last_ts = extract_first_last_timestamp(current_chunk)
        if first_ts and last_ts:
            combined_ts = combine_timestamps(first_ts, last_ts)
            
            # Filter out timestamp lines and keep only text
            text_lines = [l for l in current_chunk if not parse_timestamp(l)]
            text = ' '.join(text_lines)
            
            chunks.append(f"{combined_ts}\n{text}")
    
    # Write the processed chunks to the output file
    if output_file.lower().endswith('.docx'):
        save_to_docx(chunks, output_file)
    else:
        # Add .docx extension if not already there
        if not output_file.lower().endswith('.docx'):
            output_file += '.docx'
        save_to_docx(chunks, output_file)
    
    print(f"Processed {len(chunks)} chunks from {input_file}")

def main():
    parser = argparse.ArgumentParser(description='Process transcript files to combine timestamps and text chunks.')
    parser.add_argument('input_file', help='Path to the input transcript file (txt or docx)')
    parser.add_argument('--output_file', help='Path to save the processed output (default: input_file_processed.docx)')
    parser.add_argument('--min_lines', type=int, default=20, help='Minimum number of lines to consider as a chunk (default: 3)')
    
    args = parser.parse_args()
    
    # Generate default output filename if not provided
    if not args.output_file:
        base_name = os.path.splitext(args.input_file)[0]
        args.output_file = f"{base_name}_processed.docx"
    
    process_transcript(args.input_file, args.output_file, args.min_lines)

if __name__ == "__main__":
    main()